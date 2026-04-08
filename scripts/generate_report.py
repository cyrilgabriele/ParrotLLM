"""Generate architecture tuning report as Word doc."""
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

doc.add_heading("ParrotLLM Architecture & HP Tuning Report", level=0)
doc.add_paragraph("Status: In Progress (HP tuning running) | Date: 2026-04-06")

# --- Section 1: Architecture Search ---
doc.add_heading("1. Architecture Search Summary", level=1)
doc.add_paragraph(
    "We performed a staged architecture search to find the optimal transformer configuration "
    "within the 40M parameter budget. The search was conducted at two proxy scales (8.75M and 17.5M) "
    "to validate that architectural preferences transfer across model sizes."
)

doc.add_heading("1.1 Stage 1: 8.75M Proxy (80 trials)", level=2)
doc.add_paragraph("Search space: d_model [104-144], n_layers [6-27], n_heads [2,4,6,8,12,16], d_ff [288-384]")
doc.add_paragraph("Setup: 3000 steps, context=256, batch_size=32, RTX 5090 + torch.compile via WSL2")

table = doc.add_table(rows=6, cols=6)
table.style = "Light Grid Accent 1"
for i, h in enumerate(["Rank", "Trial", "PPL", "d_model", "n_layers", "n_heads"]):
    table.rows[0].cells[i].text = h
for r, row in enumerate([
    ["1", "25", "155.5", "144", "7", "2"],
    ["2", "15", "155.6", "144", "7", "2"],
    ["3", "58", "155.7", "136", "8", "2"],
    ["4", "73", "156.8", "144", "7", "2"],
    ["5", "33", "156.9", "136", "9", "2"],
]):
    for c, val in enumerate(row):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Key finding: ")
run.bold = True
p.add_run(
    "Wide & shallow (d_model=144, 7 layers) with 2 attention heads dominates. "
    "d_ff/d_model ratio of 2.0 beats the standard 2.67."
)

doc.add_heading("1.2 Stage 2: 17.5M Validation (40 trials)", level=2)
doc.add_paragraph("Narrowed search space based on 8.75M results. d_model [208-256], n_layers [6-14], n_heads [2,3,4,6]")

table2 = doc.add_table(rows=6, cols=7)
table2.style = "Light Grid Accent 1"
for i, h in enumerate(["Rank", "Trial", "PPL", "d_model", "n_layers", "n_heads", "d_head"]):
    table2.rows[0].cells[i].text = h
for r, row in enumerate([
    ["1", "32", "129.4", "256", "7", "4", "64"],
    ["2", "29", "129.8", "256", "7", "4", "64"],
    ["3", "3", "130.3", "256", "7", "4", "64"],
    ["4", "37", "130.4", "256", "8", "2", "128"],
    ["5", "1", "130.7", "240", "10", "3", "80"],
]):
    for c, val in enumerate(row):
        table2.rows[r + 1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Key finding: ")
run.bold = True
p.add_run(
    "Pattern confirmed at 2x scale. d_model=256, 7 layers, 4 heads (d_head=64) wins. "
    "d_head=64 aligns with best practice (GPT-3, LLaMA)."
)

# --- Section 2: Architecture Comparison ---
doc.add_heading("2. Architecture Comparison: Old vs New", level=1)

table3 = doc.add_table(rows=10, cols=4)
table3.style = "Light Grid Accent 1"
for i, h in enumerate(["Parameter", "Old (MobileLLM)", "New (Tuned)", "Change"]):
    table3.rows[0].cells[i].text = h
for r, row in enumerate([
    ["d_model", "320", "384", "+20%"],
    ["n_layers", "16", "14", "-12%"],
    ["n_heads", "8", "6", "-2"],
    ["d_head", "40", "64", "+60%"],
    ["d_ff", "854", "768", "-10%"],
    ["d_ff/d_model", "2.67", "2.00", "-0.67"],
    ["Total params", "35.8M", "39.96M", "+4.2M"],
    ["Embedding params", "16.1M (45%)", "19.3M (48%)", "+3.2M"],
    ["Transformer params", "19.7M (55%)", "20.7M (52%)", "+1.0M"],
]):
    for c, val in enumerate(row):
        table3.rows[r + 1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Concern: ")
run.bold = True
p.add_run(
    "77% of the extra 4.2M params went to embeddings (vocab=50258 x wider d_model). "
    "Total FFN capacity actually decreased (-731K). The old architecture has more "
    "learning capacity despite fewer total params."
)

# --- Section 3: HP Tuning Comparison ---
doc.add_heading("3. Hyperparameter Tuning: Old vs New Architecture", level=1)
doc.add_paragraph("Both studies use identical proxy setup: 3000 steps, context=256, same data.")

doc.add_heading("3.1 Old Architecture (d=320, 16L, 8H) - 20 trials, 6 complete", level=2)
table4 = doc.add_table(rows=4, cols=7)
table4.style = "Light Grid Accent 1"
for i, h in enumerate(["Rank", "Trial", "PPL", "lr", "dropout", "schedule", "eff. batch"]):
    table4.rows[0].cells[i].text = h
for r, row in enumerate([
    ["1", "8", "68.0", "4.26e-4", "0.015", "wsd", "256"],
    ["2", "19", "69.7", "4.53e-4", "0.027", "cosine", "256"],
    ["3", "4", "70.6", "1.31e-3", "0.073", "cosine", "128"],
]):
    for c, val in enumerate(row):
        table4.rows[r + 1].cells[c].text = val

doc.add_heading("3.2 New Architecture (d=384, 14L, 6H) - 11/80 trials, 4 complete (IN PROGRESS)", level=2)
table5 = doc.add_table(rows=4, cols=7)
table5.style = "Light Grid Accent 1"
for i, h in enumerate(["Rank", "Trial", "PPL", "lr", "dropout", "schedule", "eff. batch"]):
    table5.rows[0].cells[i].text = h
for r, row in enumerate([
    ["1", "4", "72.3", "2.14e-4", "0.084", "wsd", "256"],
    ["2", "8", "79.0", "1.06e-3", "0.111", "wsd", "64"],
    ["3", "2", "101.5", "3.07e-4", "0.043", "wsd", "64"],
]):
    for c, val in enumerate(row):
        table5.rows[r + 1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Current gap: ")
run.bold = True
p.add_run(
    "Old best = 68.0, New best = 72.3 (+6.3%). "
    "However, only 4 trials completed for new architecture vs 6 for old. "
    "Old study found its best at 4th completed trial. Study still running."
)

# --- Section 4: Key Insights ---
doc.add_heading("4. Key Insights So Far", level=1)
for insight in [
    "Wide & shallow consistently wins at proxy scale (8.75M and 17.5M)",
    "d_head=64 confirmed as optimal (matches GPT-3/LLaMA best practice)",
    "d_ff/d_model=2.0 beats standard SwiGLU ratio of 2.67",
    "WSD schedule beats cosine in both old and new studies",
    "Effective batch size 256 (bs=64, ga=4) consistently produces best results",
    "Large vocab (50258) penalizes wider models due to embedding cost",
    "New architecture has not yet beaten old architecture in proxy HP tuning (72.3 vs 68.0)",
]:
    doc.add_paragraph(insight, style="List Bullet")

# --- Section 5: Next Steps ---
doc.add_heading("5. Open Questions & Next Steps", level=1)
for step in [
    "Let HP tuning run to completion (80 trials) - may still beat 68.0",
    "Consider hybrid architecture: keep d_model=320 (cheap embeddings) but fix d_head to 64 (n_heads=5)",
    "If new architecture does not beat old in proxy, compare in full training (100K steps, context=1024)",
    "The proxy (3000 steps, ctx=256) may not fully reflect full training behavior",
]:
    doc.add_paragraph(step, style="List Bullet")

doc.save("C:/Users/chris/source/repos/ParrotLLM/results/architecture_tuning_report.docx")
print("Saved to results/architecture_tuning_report.docx")
